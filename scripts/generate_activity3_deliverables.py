from __future__ import annotations

from datetime import date, timedelta
from pathlib import Path
from textwrap import fill

import matplotlib.dates as mdates
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages
from matplotlib.patches import Patch
from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables" / "activity3"
OUT.mkdir(parents=True, exist_ok=True)

LOGO = ROOT / "assets" / "images" / "branding" / "netbite-menu-logo-mobile.png"
DOCX_PATH = OUT / "Week3_Activity_Lazatin_Allen.docx"
GANTT_PNG_PATH = OUT / "NetBite_Gantt_Chart.png"
GANTT_PDF_PATH = OUT / "NetBite_Gantt_Chart.pdf"
ACTIVITY_PDF_PATH = OUT / "NetBite_Activity3_Full_Report.pdf"
WEEK3_PDF_PATH = OUT / "Week3_Activity_Lazatin_Allen.pdf"
PROJECT_SECTION = "ITE-231"
PREPARED_DATE = "7/29/2026"

COLORS = {
    "ink": "#262328",
    "muted": "#676169",
    "red": "#C9484D",
    "orange": "#D18B5A",
    "sage": "#6F9489",
    "blue": "#6689A3",
    "plum": "#34212C",
    "paper": "#FAF8F6",
    "line": "#D5CED2",
    "light": "#EEE9EC",
    "white": "#FFFFFF",
}

PHASE_COLORS = {
    "Initiation & Planning": COLORS["red"],
    "System & Learning Design": COLORS["orange"],
    "Implementation": COLORS["sage"],
    "Testing & Quality Assurance": COLORS["blue"],
    "Release & Presentation": "#8E77A0",
}


def d(value: str) -> date:
    year, month, day = map(int, value.split("-"))
    return date(year, month, day)


TASKS = [
    ("Initiation & Planning", "1.1 Define problem, users, and project concept", d("2026-07-13"), d("2026-07-17"), "Completed", "Approved concept brief"),
    ("Initiation & Planning", "1.2 Finalize proposal, objectives, and scope boundaries", d("2026-07-18"), d("2026-07-24"), "Completed", "Proposal baseline"),
    ("Initiation & Planning", "1.3 Gather requirements and establish curriculum baseline", d("2026-07-20"), d("2026-07-31"), "In progress", "Requirements and curriculum map"),
    ("Initiation & Planning", "1.4 Prepare timeline, milestones, and risk plan", d("2026-07-27"), d("2026-08-02"), "Planned", "Week 3 planning package"),
    ("System & Learning Design", "2.1 Define architecture, data models, and persistence", d("2026-07-27"), d("2026-08-07"), "Planned", "Architecture and data design"),
    ("System & Learning Design", "2.2 Finalize mobile UX, navigation, and visual system", d("2026-07-27"), d("2026-08-09"), "Planned", "UI/UX design baseline"),
    ("System & Learning Design", "2.3 Validate lesson sequence and technical sources", d("2026-08-03"), d("2026-08-28"), "Planned", "Reviewed instructional content"),
    ("Implementation", "3.1 Build application shell, state, and progress tracking", d("2026-08-03"), d("2026-08-14"), "Planned", "Navigable mobile application"),
    ("Implementation", "3.2 Implement Chapters 1–3 and focused practices", d("2026-08-10"), d("2026-08-28"), "Planned", "Networking and switching module"),
    ("Implementation", "3.3 Implement Chapters 4–7 and practices", d("2026-08-24"), d("2026-09-18"), "Planned", "IPv4, subnetting, gateways, and ARP"),
    ("Implementation", "3.4 Implement Chapters 8–11 and educational CLI", d("2026-09-07"), d("2026-10-02"), "Planned", "Diagnostics, routing, VLAN, and models"),
    ("Implementation", "3.5 Implement bounded Network Sandbox", d("2026-09-21"), d("2026-10-16"), "Planned", "Interactive deterministic sandbox"),
    ("Implementation", "3.6 Complete responsive visuals and accessibility pass", d("2026-09-14"), d("2026-10-23"), "Planned", "Readable cross-device interface"),
    ("Testing & Quality Assurance", "4.1 Unit and integration testing", d("2026-09-28"), d("2026-10-30"), "Planned", "Automated test evidence"),
    ("Testing & Quality Assurance", "4.2 Mobile usability and device testing", d("2026-10-12"), d("2026-11-06"), "Planned", "Usability findings and test report"),
    ("Testing & Quality Assurance", "4.3 Correct defects and stabilize release candidate", d("2026-10-26"), d("2026-11-13"), "Planned", "Release candidate"),
    ("Release & Presentation", "5.1 Complete documentation and presentation materials", d("2026-11-02"), d("2026-11-13"), "Planned", "User, technical, and presentation files"),
    ("Release & Presentation", "5.2 Final build, submission, and presentation", d("2026-11-09"), d("2026-11-20"), "Planned", "Final NetBite package"),
]

MILESTONES = [
    ("M1", "Project proposal baseline approved", "Approved proposal and defined scope", d("2026-07-24")),
    ("M2", "Requirements and curriculum baseline completed", "Requirements list and 11-chapter curriculum map", d("2026-07-31")),
    ("M3", "System and mobile UX design completed", "Architecture, data model, navigation, and design system", d("2026-08-09")),
    ("M4", "Core application shell completed", "Navigable application with persisted progress", d("2026-08-14")),
    ("M5", "Chapters 1–3 prototype completed", "Lessons, practices, quizzes, and flashcards", d("2026-08-28")),
    ("M6", "Chapters 4–7 completed", "IPv4 through ARP learning sequence", d("2026-09-18")),
    ("M7", "Chapters 8–11 and CLI completed", "Diagnostics, routing, VLAN, and model activities", d("2026-10-02")),
    ("M8", "Network Sandbox feature completed", "Bounded topology, configuration, and deterministic tests", d("2026-10-16")),
    ("M9", "Responsive and accessibility pass completed", "Cross-device readable interface and accessible controls", d("2026-10-23")),
    ("M10", "System testing completed", "Automated and manual test report", d("2026-10-30")),
    ("M11", "Release candidate approved", "Stable build with priority defects resolved", d("2026-11-13")),
    ("M12", "Final system and presentation completed", "Application build, documentation, and presentation", d("2026-11-20")),
]

RISKS = [
    ("R1", "Scope expansion beyond the agreed beginner networking curriculum", "High", "High", 9, "Project lead", "New features are requested without removing or rescheduling existing work.", "Freeze the first-release scope; place DHCP, DNS, NAT, ACLs, STP, and advanced simulation in a future backlog.", "Obtain adviser approval, protect the core chapters, and move nonessential work after final submission."),
    ("R2", "Development tasks take longer than estimated", "Medium", "High", 6, "Development lead", "A milestone is more than three working days behind.", "Use weekly checkpoints, small components, reusable practice shells, and automated tests.", "Prioritize the learning path, aligned labs, and stable mobile build; defer cosmetic enhancements."),
    ("R3", "Networking content or simulation behavior is technically inaccurate", "Medium", "High", 6, "Content/QA lead", "A lesson, diagram, or simulated outcome conflicts with an RFC, IEEE reference, or validated engine.", "Record authoritative sources and test deterministic networking calculations independently of the UI.", "Correct the claim or model immediately, document the simplification, and repeat affected tests."),
    ("R4", "The interface becomes difficult to read or operate on phones/tablets", "Medium", "High", 6, "UI/UX lead", "Text clips, controls overflow, or a learner cannot complete a task at supported widths.", "Use centralized typography, responsive reflow, 44-point controls, and device-width test cases.", "Recompose dense rows vertically and simplify controls without reducing required information."),
    ("R5", "Expo or dependency updates break builds", "Medium", "High", 6, "Development lead", "Expo Doctor fails, Metro cannot bundle, or a production export stops working.", "Pin SDK-compatible versions, commit lockfiles, and run Doctor, TypeScript, lint, tests, and exports regularly.", "Return to the last validated lockfile, isolate the package change, and apply the SDK-supported version."),
    ("R6", "Saved progress or sandbox state becomes incompatible", "Low", "High", 3, "Development/QA", "A schema change causes lost progress, invalid devices, or application crashes during hydration.", "Use versioned migrations, serializable state, validation, and recovery defaults.", "Preserve compatible records, reset only corrupted state after confirmation, and document the migration."),
    ("R7", "A team member becomes unavailable", "Medium", "Medium", 4, "Project lead", "An assigned owner misses two checkpoints or reports unavailability.", "Keep tasks, interfaces, sources, and handover notes documented in the shared repository.", "Reassign the critical task, reduce parallel nonessential work, and revise target dates transparently."),
    ("R8", "Limited access to physical mobile devices delays usability testing", "Medium", "Medium", 4, "QA lead", "Required screen sizes or platforms cannot be tested during the scheduled week.", "Use Expo Go/development builds, responsive web preview, and multiple available personal devices.", "Test the highest-risk flows on available hardware first and schedule borrowed-device testing before release candidate approval."),
    ("R9", "Project files, assets, or documentation are lost or corrupted", "Low", "High", 3, "All members", "A local file becomes unavailable or differs from the shared repository without a recoverable copy.", "Use version control, remote backups, named deliverable folders, and regular validation.", "Restore the latest verified copy and reconstruct only the smallest affected change."),
    ("R10", "Generated or third-party visuals create licensing or factual problems", "Low", "Medium", 2, "Design/content", "An asset contains copied branding, embedded technical text, or inaccurate protocol details.", "Use original assets, transparent source tracking, and code-rendered technical labels.", "Replace the asset with a code-native diagram or a newly generated unlabeled object."),
]


def duration_days(start: date, end: date) -> int:
    return (end - start).days + 1


def make_gantt() -> None:
    fig, ax = plt.subplots(figsize=(16, 10.5), dpi=180)
    fig.patch.set_facecolor(COLORS["paper"])
    ax.set_facecolor(COLORS["paper"])

    y_positions = list(range(len(TASKS)))[::-1]
    for y, (phase, task, start, end, status, _) in zip(y_positions, TASKS):
        width = (end - start).days + 1
        ax.barh(
            y,
            width,
            left=mdates.date2num(start),
            height=0.62,
            color=PHASE_COLORS[phase],
            edgecolor=COLORS["ink"],
            linewidth=0.4,
            alpha=0.98,
        )
        if status == "Completed":
            ax.text(
                mdates.date2num(end) + 0.5,
                y,
                "✓",
                va="center",
                ha="left",
                fontsize=9,
                color=COLORS["sage"],
                fontweight="bold",
            )

    ax.set_yticks(y_positions)
    ax.set_yticklabels([task for _, task, *_ in TASKS], fontsize=8.3, color=COLORS["ink"])
    ax.set_xlim(mdates.date2num(d("2026-07-13")), mdates.date2num(d("2026-11-22")))
    ax.xaxis.set_major_locator(mdates.WeekdayLocator(byweekday=mdates.MO, interval=1))
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%b %d"))
    ax.tick_params(axis="x", labelsize=8, colors=COLORS["muted"], rotation=45)
    ax.grid(axis="x", color=COLORS["line"], linewidth=0.7, linestyle="-", alpha=0.9)
    ax.grid(axis="y", visible=False)
    for spine in ax.spines.values():
        spine.set_visible(False)

    ax.set_title("NetBite Project Timeline — July to November 2026", fontsize=18, fontweight="bold", color=COLORS["ink"], loc="left", pad=20)
    ax.text(
        0,
        1.015,
        "Mobile networking education game • baseline schedule • calendar days shown",
        transform=ax.transAxes,
        fontsize=9.5,
        color=COLORS["muted"],
        ha="left",
    )
    legend = [Patch(facecolor=color, edgecolor="none", label=phase) for phase, color in PHASE_COLORS.items()]
    ax.legend(
        handles=legend,
        loc="lower center",
        bbox_to_anchor=(0.5, -0.19),
        ncol=3,
        frameon=False,
        fontsize=8.5,
    )
    plt.subplots_adjust(left=0.36, right=0.98, top=0.91, bottom=0.18)
    fig.savefig(GANTT_PNG_PATH, facecolor=fig.get_facecolor())
    fig.savefig(GANTT_PDF_PATH, facecolor=fig.get_facecolor())
    plt.close(fig)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.replace("#", ""))


def set_cell_border(cell, color="D5CED2", size="6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_cell_text(cell, text: str, *, bold=False, color="262328", size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, header_fill="34212C", widths=None, font_size=8.5) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            set_cell_border(cell)
            if widths and cell_index < len(widths):
                cell.width = Inches(widths[cell_index])
            if row_index == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.name = "Aptos"
                        run.font.size = Pt(font_size)
                        run.font.color.rgb = RGBColor(255, 255, 255)
            elif row_index % 2 == 0:
                set_cell_shading(cell, "F2EEF0")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Aptos"
                    run.font.size = Pt(font_size)


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("262328")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Title", 28, "262328"),
        ("Heading 1", 19, "34212C"),
        ("Heading 2", 13.5, "C9484D"),
        ("Heading 3", 11, "D18B5A"),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("NETBITE  •  ACTIVITY 3    ")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("676169")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_header_footer(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.text = "NETBITE  /  TIMELINE & DEVELOPMENT PLAN"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string("C9484D")
    add_page_number(section.footer.paragraphs[0])


def add_section_title(doc, number: str, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(number)
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("C9484D")
    heading = doc.add_heading(title, level=1)
    heading.paragraph_format.space_before = Pt(0)
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.style = doc.styles["Subtitle"]
        p.paragraph_format.space_after = Pt(10)


def add_info_box(doc, title: str, body: str, color="EEE9EC") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    set_cell_border(cell, color="C7BEC3", size="8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title.upper())
    r.bold = True
    r.font.name = "Aptos"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("C9484D")
    p = cell.add_paragraph(body)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string("262328")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def make_document() -> None:
    doc = Document()
    set_document_defaults(doc)
    for section in doc.sections:
        add_header_footer(section)

    # Cover
    doc.sections[0].different_first_page_header_footer = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(22)
    p.add_run().add_picture(str(LOGO), width=Inches(1.25))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NETBITE")
    r.font.name = "Aptos Display"
    r.font.bold = True
    r.font.size = Pt(31)
    r.font.color.rgb = RGBColor.from_string("262328")
    p = doc.add_paragraph("Timeline & Development Plan")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style = doc.styles["Title"]
    p = doc.add_paragraph("Activity 3  •  Week 3")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = "Aptos"
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor.from_string("C9484D")

    doc.add_paragraph()
    cover_table = doc.add_table(rows=6, cols=2)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_rows = [
        ("Project", "NetBite — Mobile Networking Education Game"),
        ("Team / Section", PROJECT_SECTION),
        ("Members", "____________________________________________"),
        ("Instructor", "____________________________________________"),
        ("Prepared on", PREPARED_DATE),
        ("Submission due", "August 3, 2026 • 12:00 PM"),
    ]
    for index, (label, value) in enumerate(cover_rows):
        set_cell_text(cover_table.cell(index, 0), label.upper(), bold=True, color="676169", size=8.5)
        set_cell_text(cover_table.cell(index, 1), value, bold=index in (0, 1, 4, 5), color="262328", size=10)
        set_cell_border(cover_table.cell(index, 0), color="D5CED2")
        set_cell_border(cover_table.cell(index, 1), color="D5CED2")
        cover_table.cell(index, 0).width = Inches(1.55)
        cover_table.cell(index, 1).width = Inches(4.95)
        if index % 2 == 0:
            set_cell_shading(cover_table.cell(index, 0), "F2EEF0")
            set_cell_shading(cover_table.cell(index, 1), "F2EEF0")

    doc.add_paragraph()
    p = doc.add_paragraph(
        "A realistic implementation baseline covering project phases, detailed activities, milestones, "
        "risk response, and the Week 3 progress snapshot."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Inches(0.75)
    p.paragraph_format.right_indent = Inches(0.75)
    for run in p.runs:
        run.font.italic = True
        run.font.color.rgb = RGBColor.from_string("676169")

    doc.add_page_break()

    # Overview
    add_section_title(doc, "00", "Plan Overview", "Purpose, planning assumptions, and delivery strategy")
    doc.add_paragraph(
        "NetBite is a mobile-first networking education game that teaches introductory networking through short lessons, "
        "guided practice, quizzes, flashcards, a bounded educational CLI, and a deterministic Network Sandbox. "
        "This plan converts that product scope into a staged implementation schedule with measurable outputs."
    )
    add_info_box(
        doc,
        "Planning baseline",
        "The calendar runs from July 13 to November 20, 2026. Dates are a realistic academic-project baseline and should be aligned with the adviser’s official final-presentation calendar if it differs.",
    )
    doc.add_heading("Project objectives", level=2)
    add_bullets(
        doc,
        [
            "Deliver a readable Android/iOS learning application using React Native and Expo SDK 57.",
            "Teach 11 beginner networking chapters through accurate, source-supported, learner-paced material.",
            "Provide focused simulations that model deterministic state without claiming live packet or vendor-device emulation.",
            "Validate usability, accessibility, persistence, and technical behavior before final release.",
        ],
    )
    doc.add_heading("Scheduling approach", level=2)
    add_bullets(
        doc,
        [
            "Requirements and architecture precede high-risk simulation work.",
            "Content validation overlaps implementation so technical errors are found before final testing.",
            "Automated testing starts during implementation rather than at the end.",
            "The final three weeks protect usability testing, defect correction, documentation, and presentation preparation.",
        ],
    )

    # Landscape Gantt section
    landscape = doc.add_section(WD_SECTION.NEW_PAGE)
    landscape.orientation = WD_ORIENT.LANDSCAPE
    landscape.page_width, landscape.page_height = landscape.page_height, landscape.page_width
    landscape.top_margin = Inches(0.5)
    landscape.bottom_margin = Inches(0.5)
    landscape.left_margin = Inches(0.5)
    landscape.right_margin = Inches(0.5)
    add_header_footer(landscape)

    add_section_title(doc, "01", "Project Timeline / Gantt Chart", "Major activities from project initiation through completion")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(GANTT_PNG_PATH), width=Inches(9.65))
    p = doc.add_paragraph(
        "Status baseline: completed and in-progress labels represent the planning snapshot prepared for Activity 3. "
        "The standalone chart is also supplied as PNG and PDF."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("676169")

    doc.add_page_break()
    doc.add_heading("Detailed activity schedule", level=2)
    schedule = doc.add_table(rows=1, cols=8)
    headers = ["WBS", "Phase", "Task / Subtask", "Start", "End", "Duration", "Status", "Expected Output"]
    for i, header in enumerate(headers):
        set_cell_text(schedule.cell(0, i), header, bold=True, color="FFFFFF", size=7.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    for phase, task, start, end, status, output in TASKS:
        row = schedule.add_row().cells
        wbs, title = task.split(" ", 1)
        values = [wbs, phase, title, start.strftime("%b %d"), end.strftime("%b %d"), f"{duration_days(start, end)} days", status, output]
        for i, value in enumerate(values):
            set_cell_text(row[i], value, size=7.2, align=WD_ALIGN_PARAGRAPH.CENTER if i in (0, 3, 4, 5, 6) else WD_ALIGN_PARAGRAPH.LEFT)
    style_table(schedule, widths=[0.45, 1.15, 2.45, 0.65, 0.65, 0.65, 0.75, 1.7], font_size=7.2)

    # Return portrait
    portrait = doc.add_section(WD_SECTION.NEW_PAGE)
    portrait.orientation = WD_ORIENT.PORTRAIT
    portrait.page_width, portrait.page_height = portrait.page_height, portrait.page_width
    portrait.top_margin = Inches(0.65)
    portrait.bottom_margin = Inches(0.65)
    portrait.left_margin = Inches(0.65)
    portrait.right_margin = Inches(0.65)
    add_header_footer(portrait)

    # Milestones
    add_section_title(doc, "02", "Milestones and Deliverables", "Decision points with verifiable outputs")
    doc.add_paragraph(
        "A milestone is complete only when its expected output can be reviewed. Finishing code without its content, validation, or documentation does not satisfy the milestone."
    )
    milestone_table = doc.add_table(rows=1, cols=4)
    for i, header in enumerate(["ID", "Milestone", "Expected Output", "Target Date"]):
        set_cell_text(milestone_table.cell(0, i), header, bold=True, color="FFFFFF", size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    for mid, milestone, output, target in MILESTONES:
        row = milestone_table.add_row().cells
        values = [mid, milestone, output, target.strftime("%B %d, %Y")]
        for i, value in enumerate(values):
            set_cell_text(row[i], value, size=8.2, align=WD_ALIGN_PARAGRAPH.CENTER if i in (0, 3) else WD_ALIGN_PARAGRAPH.LEFT)
    style_table(milestone_table, widths=[0.55, 2.2, 3.1, 1.2], font_size=8.2)

    doc.add_heading("Acceptance checkpoints", level=2)
    add_bullets(
        doc,
        [
            "Content checkpoint: lesson claims and simulations agree with the documented networking sources and stated scope boundaries.",
            "Experience checkpoint: priority flows work at supported phone/tablet widths with readable text and accessible controls.",
            "Engineering checkpoint: automated tests, TypeScript, lint, Expo Doctor, asset checks, and production exports pass.",
            "Release checkpoint: documentation, build instructions, presentation materials, and final application package are complete.",
        ],
    )

    # Risk section landscape
    risk_section = doc.add_section(WD_SECTION.NEW_PAGE)
    risk_section.orientation = WD_ORIENT.LANDSCAPE
    risk_section.page_width, risk_section.page_height = risk_section.page_height, risk_section.page_width
    risk_section.top_margin = Inches(0.5)
    risk_section.bottom_margin = Inches(0.5)
    risk_section.left_margin = Inches(0.45)
    risk_section.right_margin = Inches(0.45)
    add_header_footer(risk_section)

    add_section_title(doc, "03", "Risk Assessment and Contingency Plan", "Ten project risks with triggers, prevention, and response")
    doc.add_paragraph(
        "Risk score uses a simple 1–3 likelihood × 1–3 impact matrix. Scores 6–9 require active monitoring; scores 3–4 require scheduled review; score 1–2 remains on the watch list."
    )
    risk_table = doc.add_table(rows=1, cols=8)
    risk_headers = ["ID", "Risk", "L", "I", "Score", "Owner / Trigger", "Preventive Action", "Contingency Plan"]
    for i, header in enumerate(risk_headers):
        set_cell_text(risk_table.cell(0, i), header, bold=True, color="FFFFFF", size=7.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    for rid, risk, likelihood, impact, score, owner, trigger, prevention, contingency in RISKS:
        row = risk_table.add_row().cells
        values = [rid, risk, likelihood, impact, str(score), f"{owner}\n\nTrigger: {trigger}", prevention, contingency]
        for i, value in enumerate(values):
            set_cell_text(row[i], value, size=6.8, align=WD_ALIGN_PARAGRAPH.CENTER if i in (0, 2, 3, 4) else WD_ALIGN_PARAGRAPH.LEFT)
        score_fill = "F4CCCC" if score >= 6 else "FCE5CD" if score >= 3 else "D9EAD3"
        set_cell_shading(row[4], score_fill)
    style_table(risk_table, widths=[0.38, 1.75, 0.35, 0.35, 0.45, 1.75, 2.45, 2.45], font_size=6.8)

    # Return portrait for Week 3 and closing
    final_portrait = doc.add_section(WD_SECTION.NEW_PAGE)
    final_portrait.orientation = WD_ORIENT.PORTRAIT
    final_portrait.page_width, final_portrait.page_height = final_portrait.page_height, final_portrait.page_width
    final_portrait.top_margin = Inches(0.65)
    final_portrait.bottom_margin = Inches(0.65)
    final_portrait.left_margin = Inches(0.7)
    final_portrait.right_margin = Inches(0.7)
    add_header_footer(final_portrait)

    add_section_title(doc, "04", "Week 3 Progress Presentation", "Presentation-ready progress record and evidence")
    add_info_box(
        doc,
        "Week 3 status",
        "The planning baseline is complete: concept, scope, technical stack, learning structure, visual direction, timeline, milestones, and risk responses are documented. The project can proceed into detailed design and staged implementation.",
        color="E5EFEC",
    )
    doc.add_heading("Completed by the end of Week 3", level=2)
    progress_items = [
        ("Project concept", "Defined NetBite as a mobile-first networking education game for beginners."),
        ("Problem and target users", "Focused on learners who find networking terminology and configuration difficult without guided visual practice."),
        ("Scope boundaries", "Included beginner Ethernet, switching, IPv4, subnetting, routing, VLANs, models, and bounded simulations; excluded live packet/hardware emulation."),
        ("Technology decision", "Selected React Native with Expo for Android/iOS development and web as a preview/testing surface."),
        ("Learning structure", "Established lessons, worked examples, checkpoints, mini labs, quizzes, flashcards, and optional sandbox exploration."),
        ("Visual direction", "Defined a restrained retro-industrial console style with readable typography and code-rendered technical diagrams."),
        ("Development plan", "Completed the implementation timeline, milestone outputs, and ten-risk contingency register."),
    ]
    progress_table = doc.add_table(rows=1, cols=2)
    set_cell_text(progress_table.cell(0, 0), "WORKSTREAM", bold=True, color="FFFFFF", size=8.5)
    set_cell_text(progress_table.cell(0, 1), "WEEK 3 OUTPUT", bold=True, color="FFFFFF", size=8.5)
    for workstream, output in progress_items:
        row = progress_table.add_row().cells
        set_cell_text(row[0], workstream, bold=True, size=8.4)
        set_cell_text(row[1], output, size=8.4)
    style_table(progress_table, widths=[1.65, 5.0], font_size=8.4)

    doc.add_heading("Evidence available for presentation", level=2)
    add_bullets(
        doc,
        [
            "Project and system-design documentation describing product purpose, architecture, scope, content rules, and visual standards.",
            "A functioning Expo/React Native repository that demonstrates the chosen technology is feasible.",
            "A structured curriculum and deterministic networking-domain approach that can be implemented and tested incrementally.",
            "Initial branded visual assets and a mobile-first interface direction suitable for prototype demonstrations.",
        ],
    )
    doc.add_heading("Issues and decisions", level=2)
    decision_table = doc.add_table(rows=1, cols=3)
    for i, header in enumerate(["Issue / Question", "Decision", "Reason"]):
        set_cell_text(decision_table.cell(0, i), header, bold=True, color="FFFFFF", size=8.3)
    decisions = [
        ("Should NetBite reproduce Cisco Packet Tracer?", "No. Use bounded deterministic simulations.", "Protects educational accuracy and achievable project scope."),
        ("Should every lesson contain a lab?", "No. Add practice only when it directly reinforces the lesson.", "Avoids decorative or weakly related activities."),
        ("Should technical facts be baked into generated artwork?", "No. Render labels and values in the application.", "Keeps diagrams exact, responsive, selectable, and maintainable."),
    ]
    for issue, decision, reason in decisions:
        row = decision_table.add_row().cells
        for i, value in enumerate([issue, decision, reason]):
            set_cell_text(row[i], value, size=8.1)
    style_table(decision_table, widths=[2.15, 2.1, 2.4], font_size=8.1)

    doc.add_heading("Week 4 priorities", level=2)
    add_bullets(
        doc,
        [
            "Confirm the implementation timeline, milestones, risk owners, and adviser calendar.",
            "Finalize architecture and mobile navigation before expanding feature work.",
            "Prioritize a vertical slice: lesson → focused practice → quiz → progress persistence.",
            "Begin automated validation early and record authoritative networking sources with each chapter.",
        ],
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("END OF ACTIVITY 3 DEVELOPMENT PLAN")
    r.bold = True
    r.font.name = "Aptos"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("C9484D")

    doc.core_properties.title = "NetBite Timeline & Development Plan"
    doc.core_properties.subject = "Activity 3 — Week 3"
    doc.core_properties.author = "NetBite Project Team"
    doc.core_properties.keywords = "NetBite, timeline, Gantt chart, milestones, risk assessment, Week 3 progress"
    doc.save(DOCX_PATH)


def add_slide_title(ax, number: str, title: str, subtitle: str | None = None) -> None:
    ax.text(0.06, 0.91, number, transform=ax.transAxes, fontsize=10, fontweight="bold", color=COLORS["red"], va="top")
    ax.text(0.06, 0.84, title, transform=ax.transAxes, fontsize=24, fontweight="bold", color=COLORS["ink"], va="top")
    if subtitle:
        ax.text(0.06, 0.76, subtitle, transform=ax.transAxes, fontsize=10.5, color=COLORS["muted"], va="top")
    ax.plot([0.06, 0.94], [0.71, 0.71], transform=ax.transAxes, color=COLORS["line"], linewidth=1)


def slide_bullets(ax, items: list[str], x=0.08, y=0.64, line_gap=0.105, size=15) -> None:
    for index, item in enumerate(items):
        yy = y - index * line_gap
        ax.text(x, yy, "■", transform=ax.transAxes, color=COLORS["orange"], fontsize=8, va="top")
        ax.text(x + 0.035, yy, item, transform=ax.transAxes, color=COLORS["ink"], fontsize=size, va="top", wrap=True)


def make_pdf_slide(number: str, title: str, subtitle: str | None = None):
    fig, ax = plt.subplots(figsize=(13.333, 7.5), dpi=150)
    fig.patch.set_facecolor(COLORS["paper"])
    ax.axis("off")
    add_slide_title(ax, number, title, subtitle)
    return fig, ax


def add_pdf_footer(ax, page: int, total: int) -> None:
    ax.text(0.06, 0.035, f"{PROJECT_SECTION}  •  {PREPARED_DATE}", transform=ax.transAxes, ha="left", fontsize=8, color=COLORS["muted"])
    ax.text(0.94, 0.035, f"NETBITE / ACTIVITY 3     {page} / {total}", transform=ax.transAxes, ha="right", fontsize=8, color=COLORS["muted"])


def add_slide_table(ax, headers, rows, widths, *, font_size=7.6, wrap_limits=None) -> None:
    if wrap_limits:
        wrapped_rows = []
        for row in rows:
            wrapped_rows.append([
                fill(str(value), width=wrap_limits[index]) if index < len(wrap_limits) and wrap_limits[index] else str(value)
                for index, value in enumerate(row)
            ])
    else:
        wrapped_rows = [[str(value) for value in row] for row in rows]
    table = ax.table(
        cellText=wrapped_rows,
        colLabels=headers,
        colWidths=widths,
        cellLoc="left",
        loc="center",
        bbox=[0.05, 0.075, 0.90, 0.60],
    )
    table.auto_set_font_size(False)
    table.set_fontsize(font_size)
    for (row, column), cell in table.get_celld().items():
        cell.set_edgecolor(COLORS["line"])
        cell.set_linewidth(0.7)
        cell.PAD = 0.045
        if row == 0:
            cell.set_facecolor(COLORS["plum"])
            cell.get_text().set_color(COLORS["white"])
            cell.get_text().set_weight("bold")
            cell.get_text().set_ha("center")
        else:
            cell.set_facecolor(COLORS["white"] if row % 2 else COLORS["light"])
            cell.get_text().set_color(COLORS["ink"])
            cell.get_text().set_va("center")


def make_activity_pdf() -> None:
    total_pages = 11
    with PdfPages(ACTIVITY_PDF_PATH) as pdf:
        # 1 — Cover
        fig, ax = plt.subplots(figsize=(13.333, 7.5), dpi=150)
        fig.patch.set_facecolor(COLORS["paper"])
        ax.axis("off")
        logo = Image.open(LOGO)
        ax.imshow(logo, extent=(0.08, 0.22, 0.58, 0.83), transform=ax.transAxes)
        ax.text(0.08, 0.49, "NETBITE", transform=ax.transAxes, fontsize=31, fontweight="bold", color=COLORS["ink"])
        ax.text(0.08, 0.38, "ACTIVITY 3 / WEEK 3", transform=ax.transAxes, fontsize=23, fontweight="bold", color=COLORS["red"])
        ax.text(0.08, 0.29, "Timeline & Development Plan", transform=ax.transAxes, fontsize=15, color=COLORS["muted"])
        ax.text(0.08, 0.14, "STUDENT", transform=ax.transAxes, fontsize=8, fontweight="bold", color=COLORS["muted"])
        ax.text(0.18, 0.14, "Allen Lazatin", transform=ax.transAxes, fontsize=11, fontweight="bold", color=COLORS["ink"])
        ax.text(0.42, 0.14, "SECTION", transform=ax.transAxes, fontsize=8, fontweight="bold", color=COLORS["muted"])
        ax.text(0.51, 0.14, PROJECT_SECTION, transform=ax.transAxes, fontsize=11, fontweight="bold", color=COLORS["ink"])
        ax.text(0.67, 0.14, "DATE", transform=ax.transAxes, fontsize=8, fontweight="bold", color=COLORS["muted"])
        ax.text(0.73, 0.14, PREPARED_DATE, transform=ax.transAxes, fontsize=11, fontweight="bold", color=COLORS["ink"])
        add_pdf_footer(ax, 1, total_pages)
        pdf.savefig(fig, facecolor=fig.get_facecolor())
        plt.close(fig)

        # 2 — Overview
        fig, ax = make_pdf_slide("00", "Plan Overview", "Purpose, scope, and delivery strategy")
        slide_bullets(
            ax,
            [
                "Build a readable Android/iOS networking education game with React Native and Expo.",
                "Deliver 11 beginner chapters through lessons, guided practice, quizzes, flashcards, CLI activities, and a bounded sandbox.",
                "Keep simulations deterministic and educational; do not claim live packet, hardware, timing, or vendor-OS emulation.",
                "Use staged implementation, source validation, automated tests, device checks, and a protected release period.",
                "Planning baseline: July 13–November 20, 2026; adjust only if the official adviser calendar differs.",
            ],
            y=0.63,
            line_gap=0.105,
            size=13.2,
        )
        add_pdf_footer(ax, 2, total_pages)
        pdf.savefig(fig, facecolor=fig.get_facecolor())
        plt.close(fig)

        # 3 — Gantt
        fig, ax = make_pdf_slide("01", "Project Timeline / Gantt Chart", "Major activities from initiation through completion")
        chart_ax = ax.inset_axes([0.04, 0.055, 0.92, 0.64])
        chart_ax.imshow(Image.open(GANTT_PNG_PATH))
        chart_ax.axis("off")
        add_pdf_footer(ax, 3, total_pages)
        pdf.savefig(fig, facecolor=fig.get_facecolor())
        plt.close(fig)

        # 4–5 — Detailed schedule
        for page_offset, task_chunk in enumerate((TASKS[:9], TASKS[9:]), start=4):
            fig, ax = make_pdf_slide(
                f"01.{page_offset - 3}",
                "Detailed Activity Schedule",
                "Tasks, dates, duration, status, and expected output",
            )
            rows = []
            for phase, task, start, end, status, output in task_chunk:
                wbs, title = task.split(" ", 1)
                rows.append([
                    wbs,
                    phase,
                    title,
                    start.strftime("%b %d"),
                    end.strftime("%b %d"),
                    f"{duration_days(start, end)}d",
                    status,
                    output,
                ])
            add_slide_table(
                ax,
                ["WBS", "Phase", "Task / Subtask", "Start", "End", "Dur.", "Status", "Output"],
                rows,
                [0.05, 0.13, 0.27, 0.07, 0.07, 0.05, 0.08, 0.18],
                font_size=6.8,
                wrap_limits=[5, 19, 35, 8, 8, 5, 10, 25],
            )
            add_pdf_footer(ax, page_offset, total_pages)
            pdf.savefig(fig, facecolor=fig.get_facecolor())
            plt.close(fig)


def make_week3_progress_pdf() -> None:
    total_pages = 6
    with PdfPages(WEEK3_PDF_PATH) as pdf:
        fig, ax = plt.subplots(figsize=(13.333, 7.5), dpi=150)
        fig.patch.set_facecolor(COLORS["paper"])
        ax.axis("off")
        logo = Image.open(LOGO)
        ax.imshow(logo, extent=(0.08, 0.22, 0.57, 0.82), transform=ax.transAxes)
        ax.text(0.08, 0.48, "NETBITE", transform=ax.transAxes, fontsize=31, fontweight="bold", color=COLORS["ink"])
        ax.text(0.08, 0.37, "WEEK 3 PROGRESS", transform=ax.transAxes, fontsize=23, fontweight="bold", color=COLORS["red"])
        ax.text(0.08, 0.28, "Mobile networking education game", transform=ax.transAxes, fontsize=14, color=COLORS["muted"])
        ax.text(
            0.08,
            0.11,
            f"STUDENT: ALLEN LAZATIN     SECTION: {PROJECT_SECTION}     DATE: {PREPARED_DATE}",
            transform=ax.transAxes,
            fontsize=10,
            fontweight="bold",
            color=COLORS["muted"],
        )
        ax.text(0.94, 0.05, "1 / 6", transform=ax.transAxes, ha="right", fontsize=8, color=COLORS["muted"])
        pdf.savefig(fig, facecolor=fig.get_facecolor())
        plt.close(fig)

        slides = [
            (
                "01",
                "Problem and Proposed Solution",
                "Why the project exists",
                [
                    "Beginner networking concepts are often dense, abstract, and difficult to practice safely.",
                    "NetBite breaks the learning path into short explanations, accurate visuals, and focused decisions.",
                    "The product supports conceptual understanding; it does not claim to replace real networking equipment or Packet Tracer.",
                ],
            ),
            (
                "02",
                "Work Completed",
                "Week 3 baseline outputs",
                [
                    "Defined the project concept, target learner, objectives, and first-release boundaries.",
                    "Selected React Native and Expo for one mobile-first Android/iOS codebase.",
                    "Established the lesson → practice → quiz → flashcard learning loop.",
                    "Defined the retro-industrial visual direction and readable mobile typography.",
                    "Completed the project timeline, milestone outputs, and risk contingency plan.",
                ],
            ),
            (
                "03",
                "Prototype and Design Evidence",
                "What can be demonstrated",
                [
                    "A functioning Expo project confirms the chosen stack and routing structure.",
                    "Project documentation records architecture, game design, curriculum, content, visuals, and simulation scope.",
                    "Initial branded assets demonstrate a coherent mobile identity.",
                    "Networking rules can be implemented as pure, testable state transitions rather than unreliable visual tricks.",
                ],
            ),
            (
                "04",
                "Key Decisions and Risks",
                "Keeping the implementation realistic",
                [
                    "Bound the simulation scope instead of reproducing a full vendor network operating system.",
                    "Add mini labs only when the activity directly reinforces the associated lesson.",
                    "Keep addresses, masks, routes, and layer labels in code-rendered text for accuracy.",
                    "Monitor scope growth, technical accuracy, mobile readability, schedule delay, and Expo compatibility.",
                ],
            ),
            (
                "05",
                "Next Steps",
                "Week 4 and the following development stage",
                [
                    "Finalize system architecture, persistence, mobile navigation, and reusable learning components.",
                    "Complete one tested vertical slice before expanding parallel feature work.",
                    "Validate content against authoritative sources and test each deterministic networking engine.",
                    "Review milestones and risks weekly, then adjust dates without sacrificing core learning quality.",
                    "Confirm the adviser’s final calendar and retain November as the planning baseline until then.",
                ],
            ),
        ]

        for slide_index, (number, title, subtitle, bullets) in enumerate(slides, start=2):
            fig, ax = make_pdf_slide(number, title, subtitle)
            slide_bullets(ax, bullets, y=0.62, line_gap=0.105 if len(bullets) >= 5 else 0.135, size=14.2)
            ax.text(0.06, 0.05, f"{PROJECT_SECTION}  •  {PREPARED_DATE}", transform=ax.transAxes, ha="left", fontsize=8, color=COLORS["muted"])
            ax.text(0.94, 0.05, f"NETBITE / WEEK 3     {slide_index} / 6", transform=ax.transAxes, ha="right", fontsize=8, color=COLORS["muted"])
            pdf.savefig(fig, facecolor=fig.get_facecolor())
            plt.close(fig)

def validate_outputs() -> None:
    from docx import Document as OpenDocument

    document = OpenDocument(DOCX_PATH)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    combined = text + "\n" + table_text

    required = [
        "Project Timeline / Gantt Chart",
        "Milestones and Deliverables",
        "Risk Assessment and Contingency Plan",
        "Week 3 Progress Presentation",
        "Final build, submission, and presentation",
    ]
    for value in required:
        if value not in combined:
            raise RuntimeError(f"Missing required document content: {value}")
    if len(document.tables) < 7:
        raise RuntimeError("Expected at least seven native Word tables.")
    if len(RISKS) < 5:
        raise RuntimeError("Activity requires at least five risks.")
    if len(MILESTONES) < 6:
        raise RuntimeError("Milestone table is unexpectedly incomplete.")

    for path in (DOCX_PATH, GANTT_PNG_PATH, GANTT_PDF_PATH, WEEK3_PDF_PATH):
        if not path.exists() or path.stat().st_size < 10_000:
            raise RuntimeError(f"Output is missing or unexpectedly small: {path}")

    print(f"Created: {DOCX_PATH}")
    print(f"Created: {GANTT_PNG_PATH}")
    print(f"Created: {GANTT_PDF_PATH}")
    print(f"Created: {WEEK3_PDF_PATH}")
    print(f"Validated Word sections={len(document.sections)}, tables={len(document.tables)}, paragraphs={len(document.paragraphs)}")


if __name__ == "__main__":
    make_gantt()
    make_document()
    make_week3_progress_pdf()
    validate_outputs()
